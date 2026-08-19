import io
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.services.grading_scale import grade_for

def set_cell_background(cell, fill_color: str):
    """Set background shading color of a table cell (e.g. '0F2942' for deep navy)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def generate_thesis_docx_report(submission, results, summary, narrative_text: str) -> io.BytesIO:
    """
    Generates a professional Word (.docx) assessment document matching the official Critical
    Assessment Report format.

    `summary` carries the aggregate mark and its Appendix 4.1 grade band; pass an empty dict to omit
    the mark table (for example when no criterion was successfully scored).
    """
    doc = Document()

    # Define standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. Document Header Banner
    degree_text = (submission.degree_level or "msc").upper()
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(2)
    r_h = header_p.add_run(f"CRITICAL ASSESSMENT REPORT ON {degree_text} THESIS")
    r_h.font.name = "Arial"
    r_h.font.size = Pt(14)
    r_h.font.bold = True
    r_h.font.color.rgb = RGBColor(15, 41, 66)

    header_sub = doc.add_paragraph()
    header_sub.paragraph_format.space_before = Pt(0)
    header_sub.paragraph_format.space_after = Pt(16)
    r_h_sub = header_sub.add_run("Supervisor’s Review and Corrective Guidance to the Supervisee")
    r_h_sub.font.name = "Arial"
    r_h_sub.font.size = Pt(11)
    r_h_sub.font.italic = True
    r_h_sub.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph()

    # Determine dynamic recommendation based on score if submission.supervisor_recommendation is None
    rec = submission.supervisor_recommendation
    if not rec:
        scored = [r for r in (results or []) if getattr(r, 'ai_score', None) is not None]
        total_obtained = sum(r.ai_score for r in scored) if scored else 0
        total_max = sum(r.sub_criterion.max_marks for r in scored if hasattr(r, 'sub_criterion') and r.sub_criterion) if scored else 0
        pct = (total_obtained / total_max * 100.0) if total_max > 0 else None
        band = grade_for(pct)
        rec = band.get("recommendation_detail", "Assessment incomplete")

    # 2. Metadata Table
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    metadata = [
        ("Candidate Name", submission.student_name or "N/A"),
        ("Student / Index Number", getattr(submission, 'index_number', None) or "N/A"),
        ("Programme", submission.programme or "Computer Science"),
        ("Institution", f"{submission.institution or 'Kwame Nkrumah University of Science and Technology, Kumasi'}"),
        ("Thesis Title", submission.title or "N/A"),
        ("Assessment Type", "Critical Supervisor Assessment"),
        ("Overall Recommendation", rec)
    ]

    for idx, (label, val) in enumerate(metadata):
        row = table.rows[idx]
        
        # Label cell
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_margins(cell_lbl, top=120, bottom=120, left=150, right=150)
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(15, 41, 66)

        # Value cell
        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        set_cell_margins(cell_val, top=120, bottom=120, left=150, right=150)
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.name = "Arial"
        r_val.font.size = Pt(9.5)
        if label == "Overall Recommendation":
            r_val.font.bold = True
            r_val.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph()



    # 4. Parse Markdown Report into Styled Docx Paragraphs and Tables
    lines = narrative_text.split("\n")
    in_table = False
    table_rows_data = []

    def flush_table():
        nonlocal in_table, table_rows_data
        if not table_rows_data:
            return
        
        # Check if first row is header
        headers = table_rows_data[0]
        data_rows = table_rows_data[1:]

        t = doc.add_table(rows=len(table_rows_data), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False

        col_widths = [Inches(0.6), Inches(1.8), Inches(2.0), Inches(2.1)]

        # Header Row
        hdr_row = t.rows[0]
        for c_idx, cell_text in enumerate(headers):
            cell = hdr_row.cells[c_idx]
            cell.width = col_widths[min(c_idx, len(col_widths)-1)]
            set_cell_background(cell, "0F2942")
            set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

        # Data Rows
        for r_idx, row_data in enumerate(data_rows):
            row = t.rows[r_idx + 1]
            bg_color = "F9FBFD" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    cell.width = col_widths[min(c_idx, len(col_widths)-1)]
                    set_cell_background(cell, bg_color)
                    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                    p = cell.paragraphs[0]
                    r = p.add_run(cell_text)
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
                    if c_idx == 0:
                        r.font.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        table_rows_data = []
        in_table = False

    for line in lines:
        line_str = line.strip()

        # Handle Markdown Table lines (| Col 1 | Col 2 |)
        if line_str.startswith("|") and line_str.endswith("|"):
            parts = [p.strip() for p in line_str.split("|")[1:-1]]
            # Ignore separator line (|---|---|)
            if all(set(p) <= set("-: ") for p in parts):
                continue
            table_rows_data.append(parts)
            in_table = True
            continue
        elif in_table:
            flush_table()

        if not line_str:
            continue

        p = doc.add_paragraph()

        if line_str.startswith("# "):
            r = p.add_run(line_str[2:])
            r.font.name = "Arial"
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 41, 66)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif line_str.startswith("## "):
            r = p.add_run(line_str[3:])
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 41, 66)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
        elif line_str.startswith("### "):
            r = p.add_run(line_str[4:])
            r.font.name = "Arial"
            r.font.size = Pt(10.5)
            r.font.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        elif line_str.startswith("- ") or line_str.startswith("* "):
            p.style = 'List Bullet'
            content = line_str[2:]

            # Bold label before colon if present (e.g. "- **Relevant problem:** text")
            if "**" in content:
                parts = content.split("**")
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    if i % 2 == 1:
                        r.font.bold = True
            elif ":" in content and len(content.split(":")[0]) < 40:
                label, text = content.split(":", 1)
                r_lbl = p.add_run(label + ":")
                r_lbl.font.name = "Calibri"
                r_lbl.font.size = Pt(10)
                r_lbl.font.bold = True
                
                r_txt = p.add_run(text)
                r_txt.font.name = "Calibri"
                r_txt.font.size = Pt(10)
            else:
                r = p.add_run(content)
                r.font.name = "Calibri"
                r.font.size = Pt(10)
        elif re.match(r'^\d+\.\s', line_str):
            # Numbered list item (e.g. "1. First, ...")
            r = p.add_run(line_str)
            r.font.name = "Calibri"
            r.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(3)
        else:
            # Normal paragraph text with inline bold handling
            if "**" in line_str:
                parts = line_str.split("**")
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    r = p.add_run(part)
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    if i % 2 == 1:
                        r.font.bold = True
            else:
                r = p.add_run(line_str)
                r.font.name = "Calibri"
                r.font.size = Pt(10)

    if in_table:
        flush_table()

    # 4. Add Signature Block at end of document if not already present in narrative_text
    if narrative_text and not ("Prepared by:" in narrative_text or "Signature:" in narrative_text):
        doc.add_paragraph()
        p_sig = doc.add_paragraph()
        r_sig = p_sig.add_run("Prepared by: Supervisor\nSignature: _____________________________________\nDate: __________________________________________")
        r_sig.font.name = "Arial"
        r_sig.font.size = Pt(9.5)
        r_sig.font.color.rgb = RGBColor(60, 60, 60)

    # Save to BytesIO stream
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
