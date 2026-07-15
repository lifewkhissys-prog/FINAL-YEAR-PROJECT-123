"""
Thesis Assessment Service — Multi-Agent Pipeline

Implements the 5-step rubric-grounded assessment pipeline:
  Step 0: Preprocessing (text extraction + chapter chunking)
  Step 1: Retriever (per-criterion chunk + example retrieval)
  Step 2: Scorer agent (per-criterion Groq call with structured output)
  Step 3: Verifier agent (per-criterion consistency check)
  Step 4: Aggregation (weighted score calculation)
  Step 5: Synthesis agent (full narrative report)
"""

import io
import json
import re
import logging
from datetime import datetime, timezone

import pdfplumber
import docx
from groq import AsyncGroq
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.database import SessionLocal
from app.models.thesis_critique import (
    RubricCriterion,
    GradedExample,
    ThesisSubmission,
    AssessmentResult,
    SubmissionStatus,
)
from app.services.embedding_service import embed_texts, embed_single, retrieve_top_k
from app.utils.errors import NotFoundError, ForbiddenError

logger = logging.getLogger(__name__)


# ─── Step 0: Preprocessing ──────────────────────────────────────────────────

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)
    elif ext in ("docx", "doc"):
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX files are accepted.")


def chunk_thesis_by_chapter(full_text: str) -> list[dict]:
    """Split thesis text into chapters based on heading patterns."""
    heading_re = re.compile(
        r"(?:^|\n)\s*"
        r"(?:"
        r"(?:CHAPTER\s+(?:\d+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN)(?:\s*[:\-—.]?\s*\w.*)?)"
        r"|(?:(?:INTRODUCTION|LITERATURE\s+REVIEW|METHODOLOGY|RESULTS?\s*(?:AND\s+)?(?:ANALYSIS|DISCUSSION)?|CONCLUSIONS?\s*(?:AND\s+)?(?:RECOMMENDATIONS?)?)(?:\s|$))"
        r"|(?:\d+\.\d*\s+[A-Z][A-Za-z\s,&]+)"
        r")",
        re.IGNORECASE | re.MULTILINE,
    )
    matches = list(heading_re.finditer(full_text))

    if not matches:
        return [{"title": "Full Thesis", "content": full_text}]

    chapters = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        title = m.group().strip()
        content = full_text[start:end].strip()
        if content and len(content) > 50:
            chapters.append({"title": title, "content": content})

    return chapters if chapters else [{"title": "Full Thesis", "content": full_text}]


# ─── Step 1: Retriever ──────────────────────────────────────────────────────

async def retrieve_relevant_chunks(
    criterion: RubricCriterion,
    chapters: list[dict],
    chapter_embeddings: list[list[float]],
    top_k: int = 3,
) -> list[dict]:
    """Retrieve the top-k thesis chapters most relevant to a rubric criterion."""
    if criterion.embedding is not None:
        criterion_emb = list(criterion.embedding)
    else:
        criterion_emb = embed_single(criterion.description)
    return retrieve_top_k(criterion_emb, chapter_embeddings, chapters, top_k=top_k)


async def retrieve_graded_examples(
    db: AsyncSession,
    criterion_id: int,
    top_k: int = 2,
) -> list[GradedExample]:
    """Retrieve graded examples for a specific criterion."""
    result = await db.execute(
        select(GradedExample)
        .where(GradedExample.criterion_id == criterion_id)
        .order_by(GradedExample.created_at.desc())
        .limit(top_k)
    )
    return list(result.scalars().all())


# ─── Step 2: Scorer Agent ───────────────────────────────────────────────────

async def score_criterion(
    client: AsyncGroq,
    criterion: RubricCriterion,
    relevant_chunks: list[dict],
    graded_examples: list[GradedExample],
) -> dict:
    """Score one rubric criterion via Groq with structured JSON output."""
    examples_text = ""
    if graded_examples:
        parts = []
        for ex in graded_examples:
            parts.append(
                f"Excerpt: {ex.excerpt}\n"
                f"Score: {ex.assigned_score}\n"
                f"Justification: {ex.justification or 'N/A'}"
            )
        examples_text = "\n---\n".join(parts)
    else:
        examples_text = "No reference examples available yet."

    chunks_text = "\n\n---\n\n".join(
        f"[{ch['title']}]\n{ch['content'][:3000]}" for ch in relevant_chunks
    )

    prompt = (
        f"You are assessing ONE dimension of a Master's/Bachelor's thesis. Do not evaluate "
        f"anything outside this dimension.\n\n"
        f"CRITERION: {criterion.name}\n"
        f"DESCRIPTION: {criterion.description}\n\n"
        f"SCORING GUIDE:\n"
        f"1 (weak): {criterion.level_1_desc}\n"
        f"3 (adequate): {criterion.level_3_desc}\n"
        f"5 (excellent): {criterion.level_5_desc}\n\n"
        f"REFERENCE EXAMPLES (previously graded by a human):\n{examples_text}\n\n"
        f"RELEVANT THESIS EXCERPTS TO EVALUATE:\n{chunks_text}\n\n"
        f"You MUST respond ONLY with a JSON object containing exactly these keys:\n"
        f"1. \"score\": an integer between 1 and 5.\n"
        f"2. \"justification\": a single string explaining the score based on the rubric. Do NOT return a nested object or list here.\n"
        f"3. \"cited_text\": a string containing verbatim quotes from the thesis excerpts supporting this score.\n\n"
        f"Respond ONLY in this JSON format."
    )

    try:
        response = await client.chat.completions.create(
            model=settings.GROQ_SCORER_MODEL,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.3,
            max_tokens=1024,
        )
        result = json.loads(response.choices[0].message.content)
        
        # Clamp score to 1-5
        score_val = result.get("score")
        if isinstance(score_val, dict):
            # Fallback if model puts a dict inside score
            score_val = score_val.get("score") or score_val.get("value") or 3
        try:
            result["score"] = max(1, min(5, int(score_val)))
        except (ValueError, TypeError):
            result["score"] = 3

        # Sanitize justification (ensure it is a string)
        just_val = result.get("justification")
        if isinstance(just_val, dict):
            just_val = " | ".join(f"{k}: {v}" for k, v in just_val.items())
        elif isinstance(just_val, list):
            just_val = "; ".join(str(x) for x in just_val)
        result["justification"] = str(just_val or "No justification provided.")

        # Sanitize cited_text (ensure it is a string)
        cite_val = result.get("cited_text")
        if isinstance(cite_val, dict):
            cite_val = " | ".join(f"{k}: {v}" for k, v in cite_val.items())
        elif isinstance(cite_val, list):
            cite_val = "; ".join(str(x) for x in cite_val)
        result["cited_text"] = str(cite_val or "")

        return result
    except Exception as e:
        logger.error(f"Scorer failed for '{criterion.name}': {e}")
        return {"score": 3, "justification": f"Scoring failed: {e}", "cited_text": ""}


# ─── Step 3: Verifier Agent ─────────────────────────────────────────────────

async def verify_score(
    client: AsyncGroq,
    criterion: RubricCriterion,
    score_result: dict,
) -> dict:
    """Verify a scoring decision for consistency using a second Groq call."""
    prompt = (
        f"You are verifying a grading decision, not re-grading.\n\n"
        f"CRITERION: {criterion.name}\n"
        f"SCORE GIVEN: {score_result['score']}\n"
        f"JUSTIFICATION GIVEN: {score_result['justification']}\n"
        f"CITED TEXT: {score_result['cited_text']}\n\n"
        f"Does the cited text actually support this score, given the scoring guide below?\n\n"
        f"SCORING GUIDE:\n"
        f"1: {criterion.level_1_desc}\n"
        f"3: {criterion.level_3_desc}\n"
        f"5: {criterion.level_5_desc}\n\n"
        f"Respond ONLY in a JSON object with keys 'verified' (boolean) and 'notes' (string)."
    )

    try:
        response = await client.chat.completions.create(
            model=settings.GROQ_VERIFIER_MODEL,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.2,
            max_tokens=512,
        )
        result = json.loads(response.choices[0].message.content)
        if "verified" not in result:
            result["verified"] = True
        if "notes" not in result:
            result["notes"] = ""
        return result
    except Exception as e:
        logger.error(f"Verifier failed for '{criterion.name}': {e}")
        return {"verified": True, "notes": f"Verification skipped due to error: {e}"}


# ─── Step 5: Synthesis Agent ────────────────────────────────────────────────

async def synthesize_report(
    client: AsyncGroq,
    submission: ThesisSubmission,
    criteria_results: list[dict],
    chapters: list[dict],
) -> str:
    """Generate the full narrative report via Groq (prose, not JSON)."""
    criteria_section = ""
    for cr in criteria_results:
        criteria_section += (
            f"\n### {cr['name']} (Score: {cr['score']}/5, Weight: {cr['weight']})\n"
            f"Justification: {cr['justification']}\n"
            f"Cited text: {cr['cited_text']}\n"
            f"Verifier: {'Passed' if cr.get('verified') else 'FLAGGED'}"
            f"{' — ' + cr.get('verifier_notes', '') if cr.get('verifier_notes') else ''}\n"
        )

    chapters_section = "\n\n".join(
        f"## {ch['title']}\n{ch['content'][:1200]}" for ch in chapters
    )

    prompt = (
        f"You are writing a full critical assessment report on a Master's/Bachelor's "
        f"thesis, in the style and depth of a supervisor's written review. You have "
        f"already scored the thesis against 7 rubric criteria below — use these as "
        f"your evidence base, do not re-score from scratch.\n\n"
        f"STUDENT: {submission.student_name or 'Unknown'}\n"
        f"THESIS TITLE: {submission.title or 'Untitled'}\n"
        f"PROGRAMME: {submission.programme or 'Not specified'}\n"
        f"INSTITUTION: {submission.institution or 'Not specified'}\n\n"
        f"CRITERION SCORES AND EVIDENCE:\n{criteria_section}\n\n"
        f"FULL THESIS TEXT BY CHAPTER:\n{chapters_section}\n\n"
        f"Write a report with EXACTLY these sections, matching this depth and tone:\n\n"
        f"1. OVERALL SUPERVISOR'S ASSESSMENT\n"
        f"   - 1 paragraph in second person addressing the student directly, summarizing "
        f"the thesis topic, its relevance, and overall judgement.\n"
        f"   - One bolded \"overall judgement\" line.\n\n"
        f"2. MAJOR STRENGTHS\n"
        f"   - 4-6 bullet points, each bolded with a short label followed by 1-2 sentences.\n\n"
        f"3. MAJOR CORRECTIONS REQUIRED\n"
        f"   - A markdown table: No. | Issue Identified | Why It Matters | Required Correction.\n\n"
        f"4. CHAPTER-BY-CHAPTER CRITICAL ASSESSMENT\n"
        f"   - One subsection per chapter with 3-5 bullet points each.\n\n"
        f"5. TECHNICAL AND METHODOLOGICAL COMMENTS\n"
        f"   - Bulleted, bolded sub-labels with specific technical comments.\n\n"
        f"6. FORMATTING, LANGUAGE, AND REFERENCING CORRECTIONS\n"
        f"   - Bulleted list of concrete corrections.\n\n"
        f"7. PRIORITY ACTION PLAN FOR THE CANDIDATE\n"
        f"   - A numbered, sequential list ordered by priority.\n\n"
        f"8. FINAL RECOMMENDATION\n"
        f"   - A closing paragraph plus a bolded \"Decision:\" line, plus a short "
        f"\"Supervisor's closing note\" addressed to the student by name.\n\n"
        f"Ground every claim in the criterion scores/evidence and thesis text. "
        f"Do not fabricate issues not present in the evidence."
    )

    try:
        response = await client.chat.completions.create(
            model=settings.GROQ_SYNTHESIS_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=4096,
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Synthesis failed: {e}")
        return f"*Report generation failed: {e}*"


# ─── Orchestrator: Full Pipeline ────────────────────────────────────────────

async def run_assessment_pipeline(submission_id: int, file_bytes: bytes, filename: str):
    """Run the complete multi-agent assessment pipeline as a background task."""
    logger.info(f"Starting assessment pipeline for submission {submission_id}")

    async with SessionLocal() as db:
        # Load submission
        result = await db.execute(
            select(ThesisSubmission).where(ThesisSubmission.id == submission_id)
        )
        submission = result.scalar_one_or_none()
        if not submission:
            logger.error(f"Submission {submission_id} not found")
            return

        submission.status = SubmissionStatus.assessing
        await db.commit()

        try:
            # Step 0: Extract & chunk
            if not submission.full_text:
                text = extract_text_from_file(file_bytes, filename)
                submission.full_text = text
                await db.commit()
            else:
                text = submission.full_text

            chapters = chunk_thesis_by_chapter(text)
            logger.info(f"Chunked thesis into {len(chapters)} chapters")

            # Embed chapter contents
            chapter_texts = [ch["content"][:2000] for ch in chapters]
            chapter_embeddings = embed_texts(chapter_texts)

            # Load all rubric criteria
            criteria_result = await db.execute(
                select(RubricCriterion).order_by(RubricCriterion.id)
            )
            criteria = list(criteria_result.scalars().all())

            if not criteria:
                logger.error("No rubric criteria found. Seed the rubric first.")
                submission.status = SubmissionStatus.completed
                submission.narrative_report = "*No rubric criteria configured. Please add criteria before assessing.*"
                await db.commit()
                return

            # Initialize Groq client
            client = AsyncGroq(api_key=settings.GROQ_API_KEY)

            criteria_summaries = []

            for criterion in criteria:
                logger.info(f"  Processing criterion: {criterion.name}")

                # Step 1: Retrieve relevant chunks
                relevant = await retrieve_relevant_chunks(
                    criterion, chapters, chapter_embeddings, top_k=3
                )

                # Retrieve graded examples
                examples = await retrieve_graded_examples(db, criterion.id, top_k=2)

                # Step 2: Score
                score_result = await score_criterion(client, criterion, relevant, examples)
                logger.info(f"    Score: {score_result['score']}/5")

                # Step 3: Verify
                verification = await verify_score(client, criterion, score_result)
                logger.info(f"    Verified: {verification['verified']}")

                # Store result
                ar = AssessmentResult(
                    submission_id=submission_id,
                    criterion_id=criterion.id,
                    ai_score=score_result["score"],
                    ai_justification=score_result["justification"],
                    cited_text=score_result.get("cited_text"),
                    verifier_passed=verification["verified"],
                    verifier_notes=verification.get("notes"),
                )
                db.add(ar)
                await db.commit()

                criteria_summaries.append({
                    "name": criterion.name,
                    "weight": criterion.weight,
                    "score": score_result["score"],
                    "justification": score_result["justification"],
                    "cited_text": score_result.get("cited_text", ""),
                    "verified": verification["verified"],
                    "verifier_notes": verification.get("notes", ""),
                })

            # Step 4: Aggregation
            weighted_score = sum(
                cr["weight"] * cr["score"] for cr in criteria_summaries
            )
            logger.info(f"  Weighted score: {weighted_score:.2f}/5.0")

            # Step 5: Synthesis
            logger.info("  Generating narrative report...")
            report = await synthesize_report(client, submission, criteria_summaries, chapters)

            submission.narrative_report = report
            submission.status = SubmissionStatus.completed
            await db.commit()
            logger.info(f"Assessment pipeline completed for submission {submission_id}")

        except Exception as e:
            logger.exception(f"Assessment pipeline failed for submission {submission_id}")
            submission.status = SubmissionStatus.completed
            submission.narrative_report = f"*Assessment pipeline encountered an error: {e}*"
            await db.commit()


# ─── Metadata extraction (fast model) ───────────────────────────────────────

async def extract_metadata_from_text(text: str) -> dict:
    """Use the fast Groq model to extract student name, title, programme from thesis text."""
    if not settings.GROQ_API_KEY:
        return {}
    try:
        client = AsyncGroq(api_key=settings.GROQ_API_KEY)
        response = await client.chat.completions.create(
            model=settings.GROQ_FAST_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Extract the following from this thesis text. "
                        "Return a JSON object with keys: student_name, title, programme, institution. "
                        "If a field is not found, use null.\n\n"
                        f"TEXT (first 3000 chars):\n{text[:3000]}"
                    ),
                }
            ],
            response_format={"type": "json_object"},
            temperature=0.1,
            max_tokens=256,
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logger.warning(f"Metadata extraction failed: {e}")
        return {}


# ─── CRUD: Submissions ──────────────────────────────────────────────────────

async def create_submission(
    db: AsyncSession,
    lecturer_id: int,
    student_name: str | None,
    title: str | None,
    programme: str | None,
    institution: str | None,
    filename: str,
    file_bytes: bytes,
    background_tasks,
) -> ThesisSubmission:
    """Create a new thesis submission and start text extraction."""
    text = extract_text_from_file(file_bytes, filename)

    sub = ThesisSubmission(
        lecturer_id=lecturer_id,
        student_name=(student_name or "").strip() or None,
        title=(title or "").strip() or None,
        programme=(programme or "").strip() or None,
        institution=(institution or "").strip() or None,
        file_path=filename,
        full_text=text,
        status=SubmissionStatus.pending,
    )
    db.add(sub)
    await db.commit()
    await db.refresh(sub)

    # If metadata fields are missing, extract from text in background
    if not sub.student_name or not sub.title:
        async def _fill_metadata(sid: int):
            async with SessionLocal() as s:
                r = await s.execute(select(ThesisSubmission).where(ThesisSubmission.id == sid))
                submission = r.scalar_one_or_none()
                if not submission:
                    return
                meta = await extract_metadata_from_text(submission.full_text)
                if not submission.student_name and meta.get("student_name"):
                    submission.student_name = meta["student_name"]
                if not submission.title and meta.get("title"):
                    submission.title = meta["title"]
                if not submission.programme and meta.get("programme"):
                    submission.programme = meta["programme"]
                if not submission.institution and meta.get("institution"):
                    submission.institution = meta["institution"]
                await s.commit()

        background_tasks.add_task(_fill_metadata, sub.id)

    return sub


async def trigger_assessment(
    db: AsyncSession,
    submission_id: int,
    lecturer_id: int,
    background_tasks,
) -> ThesisSubmission:
    """Trigger the assessment pipeline for a submission."""
    result = await db.execute(
        select(ThesisSubmission).where(ThesisSubmission.id == submission_id)
    )
    sub = result.scalar_one_or_none()
    if not sub:
        raise NotFoundError("Submission not found")
    if sub.lecturer_id != lecturer_id:
        raise ForbiddenError("You do not own this submission")

    # Clear previous results if re-assessing
    from sqlalchemy import delete
    await db.execute(
        delete(AssessmentResult).where(AssessmentResult.submission_id == submission_id)
    )
    sub.status = SubmissionStatus.pending
    sub.narrative_report = None
    await db.commit()

    background_tasks.add_task(
        run_assessment_pipeline,
        submission_id=sub.id,
        file_bytes=sub.full_text.encode("utf-8"),  # text already extracted
        filename=sub.file_path,
    )

    return sub


async def get_submission(db: AsyncSession, submission_id: int, lecturer_id: int) -> ThesisSubmission:
    result = await db.execute(
        select(ThesisSubmission).where(ThesisSubmission.id == submission_id)
    )
    sub = result.scalar_one_or_none()
    if not sub:
        raise NotFoundError("Submission not found")
    if sub.lecturer_id != lecturer_id:
        raise ForbiddenError("You do not own this submission")
    return sub


async def list_submissions(db: AsyncSession, lecturer_id: int) -> list[ThesisSubmission]:
    result = await db.execute(
        select(ThesisSubmission)
        .where(ThesisSubmission.lecturer_id == lecturer_id)
        .order_by(ThesisSubmission.submitted_at.desc())
    )
    return list(result.scalars().all())


async def delete_submission(db: AsyncSession, submission_id: int, lecturer_id: int) -> None:
    sub = await get_submission(db, submission_id, lecturer_id)
    await db.delete(sub)
    await db.commit()


# ─── CRUD: Assessment Results ────────────────────────────────────────────────

async def get_results(db: AsyncSession, submission_id: int, lecturer_id: int) -> list[dict]:
    """Get all criterion scores + justifications for a submission."""
    # Verify ownership
    await get_submission(db, submission_id, lecturer_id)

    result = await db.execute(
        select(AssessmentResult, RubricCriterion)
        .join(RubricCriterion, AssessmentResult.criterion_id == RubricCriterion.id)
        .where(AssessmentResult.submission_id == submission_id)
        .order_by(RubricCriterion.weight.desc())
    )
    rows = result.all()

    return [
        {
            "id": ar.id,
            "submissionId": ar.submission_id,
            "criterionId": ar.criterion_id,
            "criterionName": rc.name,
            "criterionWeight": rc.weight,
            "aiScore": ar.ai_score,
            "aiJustification": ar.ai_justification,
            "citedText": ar.cited_text,
            "verifierPassed": ar.verifier_passed,
            "verifierNotes": ar.verifier_notes,
            "supervisorOverrideScore": ar.supervisor_override_score,
            "supervisorNotes": ar.supervisor_notes,
            "createdAt": ar.created_at.isoformat() if ar.created_at else None,
        }
        for ar, rc in rows
    ]


async def override_result(
    db: AsyncSession,
    submission_id: int,
    criterion_id: int,
    lecturer_id: int,
    override_score: int,
    notes: str | None,
) -> dict:
    """Supervisor overrides a criterion score."""
    await get_submission(db, submission_id, lecturer_id)

    result = await db.execute(
        select(AssessmentResult)
        .where(
            AssessmentResult.submission_id == submission_id,
            AssessmentResult.criterion_id == criterion_id,
        )
    )
    ar = result.scalar_one_or_none()
    if not ar:
        raise NotFoundError("Assessment result not found for this criterion")

    ar.supervisor_override_score = max(1, min(5, override_score))
    ar.supervisor_notes = notes
    await db.commit()
    await db.refresh(ar)

    return {
        "id": ar.id,
        "supervisorOverrideScore": ar.supervisor_override_score,
        "supervisorNotes": ar.supervisor_notes,
    }


# ─── CRUD: Narrative Report ─────────────────────────────────────────────────

async def get_report(db: AsyncSession, submission_id: int, lecturer_id: int) -> dict:
    sub = await get_submission(db, submission_id, lecturer_id)
    return {
        "narrativeReport": sub.narrative_report,
        "narrativeReportEdited": sub.narrative_report_edited,
    }


async def update_report(
    db: AsyncSession,
    submission_id: int,
    lecturer_id: int,
    edited_text: str,
) -> dict:
    sub = await get_submission(db, submission_id, lecturer_id)
    sub.narrative_report_edited = edited_text
    sub.status = SubmissionStatus.reviewed
    await db.commit()
    return {
        "narrativeReportEdited": sub.narrative_report_edited,
        "status": sub.status.value,
    }


# ─── CRUD: Rubric Criteria ──────────────────────────────────────────────────

async def create_criterion(db: AsyncSession, data: dict) -> RubricCriterion:
    embedding = embed_single(data["description"])
    criterion = RubricCriterion(
        name=data["name"],
        description=data["description"],
        weight=data["weight"],
        level_1_desc=data["level_1_desc"],
        level_3_desc=data["level_3_desc"],
        level_5_desc=data["level_5_desc"],
        embedding=embedding,
    )
    db.add(criterion)
    await db.commit()
    await db.refresh(criterion)
    return criterion


async def list_criteria(db: AsyncSession) -> list[RubricCriterion]:
    result = await db.execute(
        select(RubricCriterion).order_by(RubricCriterion.weight.desc())
    )
    return list(result.scalars().all())


async def update_criterion(db: AsyncSession, criterion_id: int, data: dict) -> RubricCriterion:
    result = await db.execute(
        select(RubricCriterion).where(RubricCriterion.id == criterion_id)
    )
    criterion = result.scalar_one_or_none()
    if not criterion:
        raise NotFoundError("Rubric criterion not found")

    for key, val in data.items():
        if val is not None and hasattr(criterion, key):
            setattr(criterion, key, val)

    # Re-embed if description changed
    if "description" in data and data["description"] is not None:
        criterion.embedding = embed_single(data["description"])

    await db.commit()
    await db.refresh(criterion)
    return criterion


# ─── CRUD: Graded Examples ──────────────────────────────────────────────────

async def create_example(db: AsyncSession, data: dict) -> GradedExample:
    # Verify criterion exists
    crit_result = await db.execute(
        select(RubricCriterion).where(RubricCriterion.id == data["criterion_id"])
    )
    if not crit_result.scalar_one_or_none():
        raise NotFoundError("Rubric criterion not found")

    embedding = embed_single(data["excerpt"])
    example = GradedExample(
        criterion_id=data["criterion_id"],
        excerpt=data["excerpt"],
        assigned_score=data["assigned_score"],
        justification=data.get("justification"),
        embedding=embedding,
    )
    db.add(example)
    await db.commit()
    await db.refresh(example)
    return example


async def list_examples(db: AsyncSession, criterion_id: int | None = None) -> list[GradedExample]:
    query = select(GradedExample).order_by(GradedExample.created_at.desc())
    if criterion_id:
        query = query.where(GradedExample.criterion_id == criterion_id)
    result = await db.execute(query)
    return list(result.scalars().all())


# ─── DOCX Export ─────────────────────────────────────────────────────────────

def generate_docx_from_narrative(submission: ThesisSubmission) -> bytes:
    """Generate a DOCX file from the narrative report."""
    import docx.shared
    import docx.enum.text

    report_text = submission.narrative_report_edited or submission.narrative_report
    if not report_text:
        raise ValueError("No report available to export.")

    doc = docx.Document()

    # Title
    title_p = doc.add_paragraph()
    run = title_p.add_run("CRITICAL ASSESSMENT REPORT ON THESIS")
    run.font.name = "Arial"
    run.font.size = docx.shared.Pt(18)
    run.font.bold = True
    run.font.color.rgb = docx.shared.RGBColor(37, 99, 235)
    title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Supervisor's Review and Corrective Guidance")
    sub_run.font.name = "Arial"
    sub_run.font.size = docx.shared.Pt(11)
    sub_run.font.italic = True
    sub_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadata table
    meta_rows = [
        ("Candidate", submission.student_name or ""),
        ("Programme", submission.programme or ""),
        ("Institution", submission.institution or ""),
        ("Thesis Title", submission.title or ""),
    ]
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    for i, (label, val) in enumerate(meta_rows):
        c0 = meta_table.rows[i].cells[0]
        c0.text = label
        for run in c0.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = docx.shared.Pt(10)
        c1 = meta_table.rows[i].cells[1]
        c1.text = val
        for run in c1.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = docx.shared.Pt(10)

    doc.add_paragraph()

    # Parse markdown-ish report into DOCX paragraphs
    for line in report_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
        elif stripped.startswith("# "):
            h = doc.add_heading(level=1)
            r = h.add_run(stripped[2:])
            r.font.name = "Arial"
            r.font.color.rgb = docx.shared.RGBColor(37, 99, 235)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            r = h.add_run(stripped[3:])
            r.font.name = "Arial"
            r.font.color.rgb = docx.shared.RGBColor(139, 92, 246)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            r = h.add_run(stripped[4:])
            r.font.name = "Arial"
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(stripped[2:])
            r.font.name = "Arial"
            r.font.size = docx.shared.Pt(10.5)
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(re.sub(r"^\d+\.\s*", "", stripped))
            r.font.name = "Arial"
            r.font.size = docx.shared.Pt(10.5)
        else:
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.name = "Arial"
            r.font.size = docx.shared.Pt(10.5)

    doc.add_paragraph()
    for label in [
        "Prepared by: Supervisor",
        "Signature: ________________________________",
        "Date: _____________________________________",
    ]:
        p = doc.add_paragraph()
        p.add_run(label).font.size = docx.shared.Pt(10.5)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


async def export_submission_docx(
    db: AsyncSession, submission_id: int, lecturer_id: int
) -> tuple[bytes, str]:
    sub = await get_submission(db, submission_id, lecturer_id)
    docx_bytes = generate_docx_from_narrative(sub)
    return docx_bytes, sub.title or "thesis_report"
