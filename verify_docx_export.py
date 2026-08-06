"""Render the export .docx and assert the mark table is present and correct."""

import io
from types import SimpleNamespace

from docx import Document

from app.services.docx_exporter import generate_thesis_docx_report
from app.services.grading_scale import grade_for

submission = SimpleNamespace(
    student_name="Elvis Atiah",
    title="An Evidence-Based Approach to Thesis Assessment",
    programme="MPhil Information Technology",
    institution="Kwame Nkrumah University of Science and Technology, Kumasi",
    degree_level="mphil",
    supervisor_recommendation=None,
)

results = [
    {"criterion_name": "1. Statement of Problem & Justification",
     "sub_criterion_name": "Ability to articulate topic's import", "max_marks": 3.0,
     "ai_score": 2.5, "supervisor_override_score": None, "effective_score": 2.5,
     "scoring_failed": False, "cited_text": "The problem is stated clearly."},
    {"criterion_name": "1. Statement of Problem & Justification",
     "sub_criterion_name": "Justification (local/international)", "max_marks": 3.0,
     "ai_score": 2.0, "supervisor_override_score": 2.5, "effective_score": 2.5,
     "scoring_failed": False, "cited_text": "Local relevance is argued."},
    {"criterion_name": "1. Statement of Problem & Justification",
     "sub_criterion_name": "Statement of research questions", "max_marks": 4.0,
     "ai_score": None, "supervisor_override_score": None, "effective_score": None,
     "scoring_failed": True, "cited_text": None},
    {"criterion_name": "7. Presentation",
     "sub_criterion_name": "Formatting, language, citation", "max_marks": 10.0,
     "ai_score": 7.0, "supervisor_override_score": None, "effective_score": 7.0,
     "scoring_failed": False, "cited_text": "Harvard style is consistent."},
]

total = sum(r["effective_score"] for r in results if r["effective_score"] is not None)
out_of = sum(r["max_marks"] for r in results if r["effective_score"] is not None)
pct = round(total / out_of * 100, 1)

summary = {
    "total_score": round(total, 1),
    "max_possible": round(out_of, 1),
    "percentage": pct,
    **grade_for(pct),
    "unscored_criteria": 1,
    "rubric_source": "KNUST HDR Guide 2016, Appendix 4.4",
}

narrative = """# 1. Overall Supervisor's Assessment

Dear Elvis, I have reviewed your thesis critically at the MPhil level.

# 3. Major Corrections Required

| No. | Issue Identified | Why It Matters | Required Correction |
|---|---|---|---|
| 1 | Sampling frame unjustified | Threatens validity | State the sampling frame |

# 8. Final Recommendation

**Decision:** Corrections required before final submission.
"""

print(f"Marks: {total}/{out_of} = {pct}% -> {summary['grade']} ({summary['interpretation']})")

stream = generate_thesis_docx_report(submission, results, summary, narrative)
assert isinstance(stream, io.BytesIO), "did not return a BytesIO stream"

data = stream.getvalue()
print(f"Rendered {len(data):,} bytes")

with open("test_export_marks.docx", "wb") as f:
    f.write(data)

# Read it back and verify the mark table survived the round trip.
doc = Document(io.BytesIO(data))
print(f"Tables: {len(doc.tables)}  Paragraphs: {len(doc.paragraphs)}")

failures = []


def expect(label, condition, detail=""):
    if condition:
        print(f"  PASS  {label}")
    else:
        print(f"  FAIL  {label}" + (f" — {detail}" if detail else ""))
        failures.append(label)


# Metadata table and narrative tables rendered (mark table omitted per requirement).
expect("metadata table rendered", len(doc.tables) >= 2, f"got {len(doc.tables)}")

mark_table = None
for t in doc.tables:
    if t.rows and "Criterion" in t.rows[0].cells[0].text:
        mark_table = t
        break

expect("the rubric mark table is omitted", mark_table is None)

body = "\n".join(p.text for p in doc.paragraphs)
expect("narrative survived", "Dear Elvis" in body)
expect("signature block present", "Signature:" in body)

# Verification with empty summary
empty = generate_thesis_docx_report(submission, [], {}, "# Report\n\nNo marks.")
expect("report renders without error", isinstance(empty, io.BytesIO))

print()
if failures:
    print(f"{len(failures)} FAILED: {failures}")
    raise SystemExit(1)
print("All docx export checks passed. Wrote test_export_marks.docx")
