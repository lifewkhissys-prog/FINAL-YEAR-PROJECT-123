import time
import requests
import io
import docx

BASE_URL = "http://localhost:8000"  # Container internal port

def test_thesis_assessment():
    print("--- STARTING NEW THESIS SYSTEM INTEGRATION TEST ---")
    session = requests.Session()

    # 1. Register/Login Lecturer
    print("1. Logging in/Registering lecturer...")
    reg_lec = session.post(f"{BASE_URL}/auth/register", json={
        "name": "Dr. Assessment",
        "email": "examiner@devlab.edu",
        "password": "password123",
        "role": "lecturer"
    })
    
    log_lec = session.post(f"{BASE_URL}/auth/login", json={
        "email": "examiner@devlab.edu",
        "password": "password123"
    })
    assert log_lec.status_code == 200
    token = log_lec.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # 2. Verify Rubric Criteria exist
    print("2. Fetching rubric criteria...")
    criteria_res = session.get(f"{BASE_URL}/rubric/criteria", headers=headers)
    assert criteria_res.status_code == 200
    criteria = criteria_res.json()
    print(f"   Found {len(criteria)} criteria.")
    assert len(criteria) == 7

    # 3. Create a mock thesis document
    print("3. Creating mock thesis document...")
    doc = docx.Document()
    doc.add_heading('Deep Neural Networks for Medical Image Segmentation', 0)
    
    doc.add_heading('Introduction', 1)
    doc.add_paragraph(
        "This research explores the application of convolutional neural networks (CNNs) "
        "and U-Net architectures for segmenting MRI brain scans. The primary objective is to "
        "detect and delineate tumor boundaries with high accuracy and speed, helping clinicians "
        "in treatment planning. We define accuracy as the dice coefficient on the test dataset. "
        "The scope is restricted to glioblastoma cases from the BraTS 2023 dataset."
    )
    
    doc.add_heading('CHAPTER 1: LITERATURE REVIEW', 1)
    doc.add_paragraph(
        "Previous works by Ronneberger et al. (2015) introduced the U-Net for biomedical "
        "image segmentation. Recent advancements have combined U-Net with attention mechanisms "
        "(Oktay et al., 2018) to focus on target structures. However, these methods suffer from "
        "high computational costs and slow inference on edge devices."
    )
    
    doc.add_heading('CHAPTER 2: METHODOLOGY', 1)
    doc.add_paragraph(
        "We propose a Light-Unet model with parameter pruning. We split the dataset into "
        "70% training, 15% validation, and 15% test. The learning rate is set to 0.001 with "
        "Adam optimiser. Dice score is computed as 2 * |A intersect B| / (|A| + |B|)."
    )

    doc.add_heading('CHAPTER 3: RESULTS AND DISCUSSION', 1)
    doc.add_paragraph(
        "The model achieves a Dice score of 0.89 on glioblastoma segmentation, compared to "
        "0.85 for the baseline U-net. The inference time is reduced by 40% with minimal "
        "accuracy loss. All quantitative metrics are consistent between our analysis and "
        "the summary tables."
    )

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    # 4. Upload Thesis
    print("4. Uploading thesis...")
    files = {'file': ('test_thesis.docx', doc_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    upload_res = session.post(
        f"{BASE_URL}/thesis-submissions",
        headers=headers,
        data={
            "studentName": "Kofi Mensah",
            "title": "Deep Neural Networks for Medical Image Segmentation",
            "programme": "M.Sc. Computer Engineering",
            "institution": "KNUST"
        },
        files=files
    )
    print(f"   Status: {upload_res.status_code}")
    assert upload_res.status_code == 201
    submission = upload_res.json()
    submission_id = submission["id"]
    print(f"   Created submission with ID: {submission_id}")

    # 5. Trigger Assessment
    print("5. Triggering assessment pipeline...")
    assess_res = session.post(f"{BASE_URL}/thesis-submissions/{submission_id}/assess", headers=headers)
    assert assess_res.status_code == 200
    print(f"   Status: {assess_res.json()['status']}")

    # 6. Poll for completion
    print("6. Polling assessment status...")
    status = "pending"
    for i in range(12):  # 60 seconds max
        time.sleep(5)
        detail_res = session.get(f"{BASE_URL}/thesis-submissions/{submission_id}", headers=headers)
        assert detail_res.status_code == 200
        status = detail_res.json()["status"]
        print(f"   [{i*5}s] Current status: {status}")
        if status in ("completed", "reviewed"):
            break
    
    assert status in ("completed", "reviewed")

    # 7. Get Results
    print("7. Fetching criterion scores...")
    results_res = session.get(f"{BASE_URL}/thesis-submissions/{submission_id}/results", headers=headers)
    assert results_res.status_code == 200
    results = results_res.json()
    print(f"   Received {len(results)} criterion results.")
    assert len(results) == 7
    for r in results:
        print(f"     - {r['criterionName']}: AI Score = {r['aiScore']}, Verifier Passed = {r['verifierPassed']}")

    # 8. Test Override
    print("8. Testing supervisor override...")
    target_criterion_id = results[0]["criterionId"]
    override_res = session.patch(
        f"{BASE_URL}/thesis-submissions/{submission_id}/results/{target_criterion_id}",
        headers=headers,
        json={
            "supervisorOverrideScore": 5,
            "supervisorNotes": "Exemplary consistency and clarity in the mathematical formulas."
        }
    )
    assert override_res.status_code == 200
    print(f"   Override response: {override_res.json()}")

    # 9. Get Narrative Report
    print("9. Fetching narrative report...")
    report_res = session.get(f"{BASE_URL}/thesis-submissions/{submission_id}/report", headers=headers)
    assert report_res.status_code == 200
    report = report_res.json()
    print(f"   Report length: {len(report.get('narrativeReport') or '')} chars")
    assert report.get("narrativeReport") is not None

    # 10. Update/edit Narrative Report
    print("10. Updating edited narrative report...")
    edit_res = session.patch(
        f"{BASE_URL}/thesis-submissions/{submission_id}/report",
        headers=headers,
        json={"narrativeReportEdited": "## EDITED ASSESSMENT REPORT\nThis is a customized supervisor review."}
    )
    assert edit_res.status_code == 200
    print(f"    Report status now: {edit_res.json()['status']}")

    # 11. Add Graded Example
    print("11. Adding a graded example...")
    example_res = session.post(
        f"{BASE_URL}/graded-examples",
        headers=headers,
        json={
            "criterionId": target_criterion_id,
            "excerpt": "Dice score is computed as 2 * |A intersect B| / (|A| + |B|).",
            "assignedScore": 5,
            "justification": "Clear and mathematically precise definition of dice score."
        }
    )
    assert example_res.status_code == 201
    print(f"    Added example ID: {example_res.json()['id']}")

    # 12. List Graded Examples
    print("12. Listing graded examples...")
    list_ex_res = session.get(f"{BASE_URL}/graded-examples", headers=headers)
    assert list_ex_res.status_code == 200
    print(f"    Total graded examples: {len(list_ex_res.json())}")

    # 13. Export Report to Word
    print("13. Exporting report as DOCX...")
    export_res = session.get(f"{BASE_URL}/thesis-submissions/{submission_id}/export", headers=headers)
    assert export_res.status_code == 200
    assert len(export_res.content) > 0
    print(f"    DOCX exported successfully. Size: {len(export_res.content)} bytes")

    print("\n--- ALL NEW THESIS SYSTEM TESTS PASSED SUCCESSFULLY! ---")

if __name__ == "__main__":
    test_thesis_assessment()
