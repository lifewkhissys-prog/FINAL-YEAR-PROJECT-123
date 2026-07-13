import time
import requests

BASE_URL = "http://localhost:8000"  # Inside container, FastAPI is on port 8000

def test_integration():
    print("--- STARTING DEVLAB INTEGRATION TEST ---")
    session = requests.Session()

    # 1. Register Lecturer
    print("1. Registering lecturer...")
    reg_lec = session.post(f"{BASE_URL}/auth/register", json={
        "name": "Dr. Smith",
        "email": "smith@devlab.edu",
        "password": "password123",
        "role": "lecturer"
    })
    print(f"   Status: {reg_lec.status_code}, Response: {reg_lec.text}")
    assert reg_lec.status_code in (201, 409)

    # 2. Register Student
    print("2. Registering student...")
    reg_std = session.post(f"{BASE_URL}/auth/register", json={
        "name": "Jane Doe",
        "email": "jane@devlab.edu",
        "password": "password123",
        "role": "student"
    })
    print(f"   Status: {reg_std.status_code}, Response: {reg_std.text}")
    assert reg_std.status_code in (201, 409)

    # 3. Login Lecturer
    print("3. Logging in lecturer...")
    log_lec = session.post(f"{BASE_URL}/auth/login", json={
        "email": "smith@devlab.edu",
        "password": "password123"
    })
    print(f"   Status: {log_lec.status_code}")
    assert log_lec.status_code == 200
    lec_token = log_lec.json()["access_token"]

    # 4. Login Student
    print("4. Logging in student...")
    log_std = session.post(f"{BASE_URL}/auth/login", json={
        "email": "jane@devlab.edu",
        "password": "password123"
    })
    print(f"   Status: {log_std.status_code}")
    assert log_std.status_code == 200
    std_token = log_std.json()["access_token"]

    # --- Lecturer Operations ---
    lec_headers = {"Authorization": f"Bearer {lec_token}"}
    
    # 5. Create Course
    print("5. Creating a course...")
    course_res = requests.post(f"{BASE_URL}/courses", headers=lec_headers, json={
        "title": "CS101: Introduction to Programming",
        "description": "Learn python basics.",
        "language": "python"
    })
    print(f"   Status: {course_res.status_code}, Response: {course_res.json()}")
    assert course_res.status_code == 201
    course_id = course_res.json()["id"]

    # 6. Enroll Student
    print("6. Enrolling student...")
    enroll_res = requests.post(f"{BASE_URL}/courses/{course_id}/students", headers=lec_headers, json={
        "email": "jane@devlab.edu"
    })
    print(f"   Status: {enroll_res.status_code}, Response: {enroll_res.json()}")
    assert enroll_res.status_code == 201

    # 7. Create Assessment
    print("7. Creating assessment...")
    assessment_res = requests.post(f"{BASE_URL}/assessments", headers=lec_headers, json={
        "courseId": course_id,
        "title": "Midterm Exam 1",
        "startsAt": "2026-07-01T00:00:00Z",
        "endsAt": "2026-07-10T00:00:00Z"  # Active window
    })
    print(f"   Status: {assessment_res.status_code}, Response: {assessment_res.json()}")
    assert assessment_res.status_code == 201
    assessment_id = assessment_res.json()["id"]

    # 8. Create Problem
    print("8. Creating problem...")
    prob_res = requests.post(f"{BASE_URL}/problems", headers=lec_headers, json={
        "assessmentId": assessment_id,
        "title": "Add Two Numbers",
        "type": "challenge",
        "language": "python",
        "content": {
            "description": "Write a python function to add two inputs.",
            "starterCode": "def solve(a, b):\n    return a + b"
        }
    })
    print(f"   Status: {prob_res.status_code}, Response: {prob_res.json()}")
    assert prob_res.status_code == 201
    problem_id = prob_res.json()["id"]

    # 9. Add Test Cases
    print("9. Adding test cases...")
    tc_res = requests.put(f"{BASE_URL}/problems/{problem_id}/test-cases", headers=lec_headers, json=[
        {"stdin": "1 2", "expectedStdout": "3", "isHidden": False, "position": 0},
        {"stdin": "5 10", "expectedStdout": "15", "isHidden": True, "position": 1}
      ])
    print(f"   Status: {tc_res.status_code}, Response: {tc_res.json()}")
    assert tc_res.status_code == 200

    # --- Student Operations ---
    std_headers = {"Authorization": f"Bearer {std_token}"}

    # 10. Fetch Student Dashboard
    print("10. Fetching student dashboard...")
    dash_res = requests.get(f"{BASE_URL}/student/dashboard", headers=std_headers)
    print(f"    Status: {dash_res.status_code}, Response: {dash_res.json()}")
    assert dash_res.status_code == 200

    # 11. Run practice/check code
    print("11. Running practice/check code...")
    run_res = requests.post(f"{BASE_URL}/submissions/run", headers=std_headers, json={
        "problem_id": problem_id,
        "code": "def solve(a, b):\n    return int(a) + int(b)",
        "language": "python"
    })
    print(f"    Status: {run_res.status_code}, Response: {run_res.json()}")
    assert run_res.status_code == 200

    # --- Thesis Critique Operations ---
    # 12. Upload Thesis (using a mock docx bytes)
    print("12. Uploading thesis...")
    import docx
    import io
    doc = docx.Document()
    doc.add_heading('Thesis Title: Deep Learning in Code Analysis', 0)
    doc.add_paragraph('Chapter 1: Introduction')
    doc.add_paragraph('The study will look at the various parameters and try to find a pattern.')
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    files = {'file': ('thesis.docx', doc_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    thesis_res = requests.post(f"{BASE_URL}/thesis-critique", headers=lec_headers, data={
        "candidateName": "Jane Student",
        "programme": "M.Sc. Computer Science",
        "thesisTitle": "Deep Learning in Code Analysis"
    }, files=files)
    print(f"    Status: {thesis_res.status_code}, Response: {thesis_res.json()}")
    assert thesis_res.status_code == 201
    critique_id = thesis_res.json()["id"]

    # 13. Wait and check polling for Thesis Critique
    print("13. Polling for thesis critique completion...")
    for _ in range(5):
        time.sleep(2)
        det_res = requests.get(f"{BASE_URL}/thesis-critique/{critique_id}", headers=lec_headers)
        status = det_res.json()["status"]
        print(f"    Polling status: {status}")
        if status in ("completed", "failed"):
            print(f"    Report parsed: {det_res.json().get('reportJson') is not None}")
            assert status == "completed"
            break

    # 14. Fetch Lecturer Dashboard
    print("14. Fetching lecturer dashboard...")
    lec_dash_res = requests.get(f"{BASE_URL}/lecturer/dashboard", headers=lec_headers)
    print(f"    Status: {lec_dash_res.status_code}, Response: {lec_dash_res.json()}")
    assert lec_dash_res.status_code == 200

    print("--- ALL TESTS PASSED SUCCESSFULLY! ---")

if __name__ == "__main__":
    test_integration()
